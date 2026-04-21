#!/usr/bin/env python3
"""
Assessment Workflow – Batch Grading
Processes directories of submissions (.txt, .docx, .pdf) for grading.

Features:
- Text extraction with quality gates
- Optional grading via FastAPI server
- Deterministic structural rules (paragraph count)
- Idempotent output with --overwrite support
- Dry-run mode for testing

Output structure (under artifacts/runs/batches/{batch_id}/):
- grading/canonical/*.json - Final grading records
- grading/intermediate/ - Reserved for future checkpoints
- reports/final/batch_report.json - Batch summary
- reports/debug/batch_rollup.jsonl - Per-file JSONL log

Extracted text goes to: artifacts/extraction_store/
"""
import argparse
import json
import logging
import re
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("grade_batch")

DEFAULT_SERVER = "http://localhost:8000"


def read_text(path: Path) -> Tuple[str, List[str], str]:
    """Extract text from .txt, .docx, or .pdf file.
    
    Returns: (text, warnings, extraction_method)
    """
    warnings: List[str] = []
    suffix = path.suffix.lower()
    
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8")
        return text, warnings, "txt_direct"
    
    if suffix == ".docx":
        if Document is None:
            raise RuntimeError("python-docx is not installed")
        doc = Document(path)
        # First, try extracting responses from two-column tables (question | response)
        responses: List[str] = []
        try:
            for table in doc.tables:
                if len(table.rows) == 0 or len(table.columns) < 2:
                    continue
                # Skip header row if it looks like a header
                start_row = 0
                first_cell_text = (table.rows[0].cells[0].text or "").strip().lower()
                if "question" in first_cell_text or "prompt" in first_cell_text:
                    start_row = 1
                for row_idx in range(start_row, len(table.rows)):
                    row = table.rows[row_idx]
                    if len(row.cells) >= 2:
                        response_cell = row.cells[1]
                        # Preserve inner paragraph breaks: convert single \n to \n\n
                        # This is necessary because cell.text joins inner paragraphs with \n only
                        response_text = (response_cell.text or "").strip()
                        # Replace single newline with double newline (except already-double ones)
                        response_text = response_text.replace("\n\n", "\x00PLACEHOLDER\x00")  # Protect existing double newlines
                        response_text = response_text.replace("\n", "\n\n")  # Convert single to double
                        response_text = response_text.replace("\x00PLACEHOLDER\x00", "\n\n")  # Restore protected ones
                        if response_text:
                            responses.append(response_text)
        except Exception:
            # Fall through to paragraph extraction on any table parsing issue
            responses = []
        if responses:
            text = "\n\n".join(responses)
            warnings.append("docx_table_extraction")
            return text, warnings, "docx_table"
        # Fallback: use paragraphs if no table responses found
        paras = [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]
        text = "\n\n".join(paras)
        return text, warnings, "docx_paragraphs"
    
    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("pypdf is not installed")
        reader = PdfReader(path)
        parts: List[str] = []
        for page in reader.pages:
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(t)
        text = "\n\n".join(parts)
        return text, warnings, "pdf_pypdf"
    
    raise RuntimeError(f"unsupported_file_type:{suffix}")


def _shared_block_patterns() -> List[str]:
    """Block patterns matching boilerplate text - same as extract_responses.py"""
    return [
        r"Note:\s+To\s+edit\s+this\s+document.*?visible\.",
        r"Week\s+1\s+Business\s+Activity:.*?submit\s+your\s+assignment!?",
        r"Answer\s+each\s+question\s+in\s+the\s+box\s+below.*?submit\s+your\s+assignment!?",
        r"\(Box\s+will\s+expand\s+as\s+you\s+type\)",
        r"1\.\s+What\s+are\s+the\s+factors\s+of\s+production\s+needed\s+by\s+a\s+surfboard\s+manufacturer\?",
        r"2\.\s+Where\s+does\s+the\s+surfboard\s+company\s+get\s+these\s+factors\s+of\s+production\?",
        r"3\.\s+Where\s+does\s+the\s+company\s+get\s+money\s+to\s+pay\s+for\s+additional\s+resources\?",
    ]


def _clean_common_boilerplate(text: str) -> Tuple[str, List[str]]:
    """Remove boilerplate patterns - same as extract_responses.py"""
    removed: List[str] = []
    cleaned = text
    for pat in _shared_block_patterns():
        new_cleaned, n = re.subn(pat, "", cleaned, flags=re.IGNORECASE | re.DOTALL)
        if n > 0:
            removed.append(f"pattern:{pat} count:{n}")
            cleaned = new_cleaned
    # Normalize spaces while preserving paragraph breaks (\n\n)
    # Split on double newline, normalize each paragraph, then rejoin
    paras = cleaned.split("\n\n")
    normalized_paras = [re.sub(r"\s+", " ", p).strip() for p in paras if p.strip()]
    cleaned = "\n\n".join(normalized_paras)
    return cleaned, removed


def _remove_instruction_relics(text: str) -> str:
    """Remove known assignment instruction boilerplate and relics.
    
    Uses the proven logic from extract_responses.py:
    Apply block patterns, then normalize whitespace while preserving paragraph breaks.
    """
    cleaned, _ = _clean_common_boilerplate(text)
    return cleaned

def _question_prompt_patterns() -> List[str]:
    """Question prompt patterns to split responses"""
    return [
        r"what\s+are\s+the\s+factors\s+of\s+production\s+needed\s+by\s+a\s+surfboard\s+manufacturer",
        r"where\s+does\s+the\s+surfboard\s+company\s+get\s+these\s+factors\s+of\s+production",
        r"where\s+does\s+the\s+company\s+get\s+money\s+to\s+pay\s+for\s+additional\s+resources",
    ]


def _split_by_prompts(raw_text: str, prompt_patterns: List[str]) -> List[str]:
    """Split text by question prompts"""
    if not prompt_patterns:
        return []
    combined = "|".join(f"({p})" for p in prompt_patterns)
    matches = list(re.finditer(combined, raw_text, flags=re.IGNORECASE | re.DOTALL))
    if not matches:
        return []
    segments: List[str] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        seg = raw_text[start:end].strip()
        if seg:
            segments.append(seg)
    return segments


def _strip_leading_bullets(text: str) -> str:
    """Remove stray bullet-like markers (e.g., '?', '•') at paragraph starts"""
    if not text:
        return text
    bullet_re = re.compile(r"^[\?\u2022\u2023\u2043\u25E6\-\u2013\u2014\*]+\s*")
    paras = text.split("\n\n")
    cleaned = [bullet_re.sub("", p.strip()) for p in paras]
    return "\n\n".join(cleaned).strip()



def clean_text(raw: str) -> str:
    """Remove boilerplate, normalize whitespace, preserve paragraph structure.
    
    Uses the complete logic from extract_responses.py:
    1. Remove block pattern boilerplate
    2. Filter line-level keywords
    3. Split by question prompts if available
    4. Reflow into paragraphs
    5. Strip bullets and normalize
    """
    # Apply block patterns first
    cleaned = raw
    block_patterns = _shared_block_patterns()
    for bp in block_patterns:
        cleaned = re.sub(bp, "", cleaned, flags=re.IGNORECASE | re.DOTALL)
    
    # Line-level keyword blocklist
    keyword_blocklist = [
        "Note: To edit this document",
        "Week 1 Business Activity: Factors in Producing a Surfboard",
        "Use the information from Chapter",
        "Answer each question in the box below",
        "assignment drop box",
        "D2L Brightspace",
    ]
    
    compiled_prompts = [re.compile(p, re.IGNORECASE) for p in _question_prompt_patterns()]
    
    def is_template_line(ln: str) -> bool:
        low = ln.lower()
        if not low:
            return False
        # Remove page numbers
        if re.match(r"^\s*page\s*\d+\s*$", low):
            return True
        # Keyword blocklist
        for kw in keyword_blocklist:
            if kw.lower() in low:
                return True
        # Question prompts
        for rx in compiled_prompts:
            if rx.search(low):
                return True
        return False
    
    # Split by question prompts
    prompt_patterns = _question_prompt_patterns()
    prompt_splits = _split_by_prompts(cleaned, prompt_patterns)
    
    # Helper to reflow lines into paragraphs
    def _reflow_lines_to_paragraphs(ls: List[str]) -> str:
        paragraphs: List[str] = []
        buf: List[str] = []
        for ln in ls:
            s = ln.strip()
            if not s:
                if buf:
                    para = " ".join(buf)
                    para = re.sub(r"\s+", " ", para).strip()
                    paragraphs.append(para)
                    buf = []
                continue
            if s.endswith("-"):
                buf.append(s[:-1])
            else:
                buf.append(s)
        if buf:
            para = " ".join(buf)
            para = re.sub(r"\s+", " ", para).strip()
            paragraphs.append(para)
        return "\n\n".join(paragraphs)
    
    # Process prompt splits if available
    if prompt_splits:
        cleaned_chunks = []
        for chunk in prompt_splits:
            chunk_lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
            chunk_lines = [ln for ln in chunk_lines if not is_template_line(ln)]
            chunk_text = _reflow_lines_to_paragraphs(chunk_lines).strip()
            chunk_text, _ = _clean_common_boilerplate(chunk_text)
            chunk_text = _strip_leading_bullets(chunk_text)
            if chunk_text:
                cleaned_chunks.append(chunk_text)
        if cleaned_chunks:
            return "\n\n".join(cleaned_chunks)
    
    # Fallback if no prompt splits
    lines = [ln.strip() for ln in cleaned.splitlines()]
    cleaned_lines = [ln for ln in lines if not is_template_line(ln)]
    out_text = _reflow_lines_to_paragraphs(cleaned_lines).strip()
    out_text, _ = _clean_common_boilerplate(out_text)
    out_text = _strip_leading_bullets(out_text)
    return out_text or ""


def count_paragraphs(cleaned_text: str) -> int:
    """Count non-empty paragraphs separated by double newline."""
    paras = [p.strip() for p in cleaned_text.split("\n\n") if p.strip()]
    return len(paras)


def apply_quality_gates(
    raw_text: str, 
    cleaned_text: str, 
    extraction_warnings: List[str], 
    min_words: int = 30, 
    min_ratio: float = 0.15
) -> Tuple[str, List[str]]:
    """Apply quality checks to determine if submission should be graded.
    
    Returns: (status, reasons)
    - status: "OK_FOR_GRADING" or "NEEDS_REVIEW"
    - reasons: List of issues found
    
    Quality Gates:
    - min_words: 30 (reject extremely short responses)
    - min_ratio: 0.15 (15% retention after boilerplate removal)
    
    Threshold History:
    - 2026-01-19: Reduced min_ratio from 0.20 to 0.15
      Reason: PDFs with heavy boilerplate (headers, instructions) were flagged
      as NEEDS_REVIEW despite containing valid 3-paragraph responses.
      Example: anon-041-raw.pdf had 1764→345 chars (19.5%) but was complete.
      New threshold allows concise responses while word count gate (30) prevents
      truly empty submissions.
    """
    reasons: List[str] = []
    word_count = len(cleaned_text.split())
    ratio = len(cleaned_text) / max(len(raw_text), 1)
    
    if word_count < min_words:
        reasons.append(f"low_word_count:{word_count}")
    if ratio < min_ratio:
        reasons.append(f"short_ratio:{ratio:.2f}")
    
    header_footer_warnings = [w for w in extraction_warnings if "header" in w or "footer" in w]
    if header_footer_warnings:
        reasons.extend(header_footer_warnings)
    
    if reasons:
        return "NEEDS_REVIEW", reasons
    return "OK_FOR_GRADING", reasons


def grade_submission(server: str, payload: Dict[str, Any], retries: int = 3) -> Dict[str, Any]:
    """Call grading API endpoint with retry logic for connection failures.
    
    Note: OpenAI API calls can take >120s on overloaded systems; timeout is 300s.
    """
    import time as time_module
    
    last_error = None
    for attempt in range(retries):
        try:
            logger.info(f"Grading attempt {attempt + 1}/{retries}: payload_size={len(json.dumps(payload))} bytes")
            resp = requests.post(
                f"{server}/tier2/feedback-suggest", 
                json=payload, 
                timeout=300  # Increased to 300s to allow OpenAI API latency
            )
            logger.info(f"Grading attempt {attempt + 1} succeeded: status={resp.status_code}")
            if resp.status_code != 200:
                raise RuntimeError(f"grade_request_failed:{resp.status_code}:{resp.text[:200]}")
            return resp.json()
        except (requests.Timeout, requests.ConnectionError, OSError, IOError, requests.RequestException) as e:
            last_error = e
            logger.warning(f"Grading attempt {attempt + 1} error: {type(e).__name__}: {str(e)[:150]}")
            if attempt < retries - 1:
                wait_time = 5 * (attempt + 1)
                logger.warning(
                    f"Retrying in {wait_time}s..."
                )
                time_module.sleep(wait_time)
            else:
                logger.error(f"Grading API request failed after {retries} attempts")
                raise RuntimeError(f"grade_request_failed_after_retries:{str(e)[:200]}") from e
        except Exception as e:
            # Other errors (JSON parsing, HTTP 500, etc.) don't retry
            logger.error(f"Grading attempt {attempt + 1} non-retryable error: {type(e).__name__}: {str(e)[:150]}")
            raise RuntimeError(f"grade_request_failed:{str(e)[:200]}") from e
    
    # Should not reach here
    raise RuntimeError(f"grade_request_failed_unknown:{last_error}")


def persist_cleaned_text(cleaned_text: str, path: Path, overwrite: bool) -> None:
    """Save extracted text to disk."""
    if path.exists() and not overwrite:
        raise RuntimeError(f"exists:{path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(cleaned_text, encoding="utf-8")


def persist_record(record: Dict[str, Any], path: Path, overwrite: bool) -> None:
    """Save grading record as JSON."""
    if path.exists() and not overwrite:
        raise RuntimeError(f"exists:{path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(record, indent=2), encoding="utf-8")


def append_rollup(record: Dict[str, Any], path: Path) -> None:
    """Append record to JSONL rollup file."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(record) + "\n")


def apply_structural_rules(grade_data: Dict[str, Any], paragraph_count: int) -> Tuple[Dict[str, Any], List[str]]:
    """Apply deterministic structural rules to grading output.
    
    If paragraph count < 3, downgrade Adherence from Meets Expectations (15) to Needs Improvement (11.25).
    Returns adjusted grade_data and list of adjustments applied.
    """
    adjustments: List[str] = []
    adjusted_data = grade_data.copy()
    
    if paragraph_count < 3:
        # Check if Adherence score was Meets Expectations (score_high >= 15)
        score_high = grade_data.get("score_high")
        if score_high is not None and score_high >= 15:
            # Downgrade to Needs Improvement (11.25)
            adjusted_data["adherence_score"] = 11.25
            adjusted_data["adherence_original"] = score_high
            adjustments.append(f"adherence_downgraded_paragraphs:{paragraph_count}")
    
    return adjusted_data, adjustments


def build_payload(cleaned_text: str, args: argparse.Namespace) -> Dict[str, Any]:
    """Build API request payload."""
    return {
        "text": cleaned_text,
        "course": args.course_id,
        "assignment_id": args.assignment_id,
        "week": args.week,
        "points_possible": args.points,
        "top_k_rubric": args.top_k_rubric,
        "top_k_feedback": args.top_k_feedback,
    }


def run_batch(args: argparse.Namespace) -> Dict[str, Any]:
    """Main batch processing function."""
    batch_root = Path(args.output_root) / args.batch_id
    grading_dir = batch_root / "grading"
    grading_intermediate_dir = grading_dir / "intermediate"
    grading_canonical_dir = grading_dir / "canonical"
    reports_dir = batch_root / "reports"
    reports_final_dir = reports_dir / "final"
    reports_debug_dir = reports_dir / "debug"
    rollup_path = reports_debug_dir / "batch_rollup.jsonl"

    if batch_root.exists() and not args.overwrite:
        logger.error(f"Batch output already exists: {batch_root}; use --overwrite to replace")
        return {"status": "failed", "reason": "batch_exists"}

    files = sorted([p for p in Path(args.input_dir).iterdir() if p.is_file()])
    if args.limit:
        files = files[: args.limit]

    summary: Dict[str, Any] = {
        "batch_id": args.batch_id,
        "total": len(files),
        "extracted": 0,
        "graded": 0,
        "failed": 0,
        "needs_review": [],
        "dry_run": 0,
        "records": [],
    }

    if not args.dry_run:
        for d in (grading_intermediate_dir, grading_canonical_dir, reports_final_dir, reports_debug_dir):
            d.mkdir(parents=True, exist_ok=True)

    for path in files:
        record: Dict[str, Any] = {
            "batch_id": args.batch_id,
            "assignment_id": args.assignment_id,
            "course_id": args.course_id,
            "week": args.week,
            "original_filename": path.name,
            "file_type": path.suffix.lower(),
            "status": "pending",
            "extraction_warnings": [],
        }
        
        try:
            # Extract and clean text
            raw_text, warnings, method = read_text(path)
            cleaned_text = clean_text(raw_text)
            record.update({
                "extraction_method": method,
                "raw_length": len(raw_text),
                "cleaned_length": len(cleaned_text),
                "extraction_warnings": warnings,
            })

            # Apply quality gates
            quality_status, reasons = apply_quality_gates(raw_text, cleaned_text, warnings)
            record["quality_status"] = quality_status
            record["quality_reasons"] = reasons

            # Dry-run mode: skip actual grading
            if args.dry_run:
                record["status"] = "dry_run"
                summary["dry_run"] += 1
                summary["records"].append(record)
                if quality_status == "NEEDS_REVIEW":
                    summary["needs_review"].append(path.name)
                continue

            # Save extracted text to global extraction store
            extraction_store = Path("artifacts/extraction_store")
            extraction_store.mkdir(parents=True, exist_ok=True)
            # Avoid duplicate _extracted suffixes
            if path.stem.endswith("_extracted"):
                cleaned_path = extraction_store / f"{path.stem}.txt"
            else:
                cleaned_path = extraction_store / f"{path.stem}_extracted.txt"
            persist_cleaned_text(cleaned_text, cleaned_path, args.overwrite)
            record["cleaned_text_path"] = str(cleaned_path)

            # Compute paragraph count for structural checks
            paragraph_count = count_paragraphs(cleaned_text)
            record["paragraph_count"] = paragraph_count

            # Mark status based on quality
            if quality_status == "NEEDS_REVIEW":
                record["status"] = "needs_review"
                summary["needs_review"].append(path.name)
            else:
                record["status"] = "ready"

            # Grade if requested and quality is acceptable
            if quality_status == "OK_FOR_GRADING" and args.grade:
                # Add throttle delay to prevent overwhelming server
                time.sleep(3)
                logger.info(f"Starting grading for {path.name}...")
                payload = build_payload(cleaned_text, args)
                grade_data = grade_submission(args.server, payload)
                
                # Apply deterministic structural rules
                grade_data, adjustments = apply_structural_rules(grade_data, paragraph_count)
                if adjustments:
                    record["structural_adjustments"] = adjustments
                
                record["grade"] = grade_data
                record["status"] = "graded"
                summary["graded"] += 1

            # Persist record outside the grading conditional
            persist_record(record, grading_canonical_dir / f"{path.stem}.json", args.overwrite)
            append_rollup(record, rollup_path)
            summary["extracted"] += 1

        except Exception as e:
            record["status"] = "failed"
            record["error"] = str(e)
            summary["failed"] += 1
            summary["records"].append(record)
            if not args.dry_run:
                try:
                    persist_record(record, grading_canonical_dir / f"{path.stem}.json", args.overwrite)
                    append_rollup(record, rollup_path)
                except Exception:
                    pass
            continue

        summary["records"].append(record)
        time.sleep(0.2)

    # Write batch report
    if not args.dry_run:
        batch_report = {
            "batch_id": args.batch_id,
            "total": summary["total"],
            "extracted": summary["extracted"],
            "graded": summary["graded"],
            "failed": summary["failed"],
            "needs_review_count": len(summary["needs_review"]),
            "needs_review_files": summary["needs_review"],
            "dry_run": summary["dry_run"],
        }
        report_path = reports_final_dir / "batch_report.json"
        report_path.write_text(json.dumps(batch_report, indent=2), encoding="utf-8")

    logger.info(
        "Batch %s total=%s extracted=%s graded=%s failed=%s needs_review=%s dry_run=%s",
        args.batch_id,
        summary["total"],
        summary["extracted"],
        summary["graded"],
        summary["failed"],
        len(summary["needs_review"]),
        summary["dry_run"],
    )
    return summary


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments."""
    ap = argparse.ArgumentParser(description="Assessment Workflow: Batch Grading runner")
    ap.add_argument("input_dir", help="Directory containing submission files (.txt/.docx/.pdf)")
    ap.add_argument("--assignment-id", dest="assignment_id", required=True)
    ap.add_argument("--course-id", dest="course_id", required=True)
    ap.add_argument("--week", dest="week", type=int, required=True)
    ap.add_argument("--batch-id", dest="batch_id", required=True)
    ap.add_argument("--output-root", dest="output_root", default="artifacts/runs/batches")
    ap.add_argument("--server", dest="server", default=DEFAULT_SERVER)
    ap.add_argument("--points", dest="points", type=float, default=40.0)
    ap.add_argument("--top-k-rubric", dest="top_k_rubric", type=int, default=6)
    ap.add_argument("--top-k-feedback", dest="top_k_feedback", type=int, default=6)
    ap.add_argument("--dry-run", dest="dry_run", action="store_true")
    ap.add_argument("--overwrite", dest="overwrite", action="store_true")
    ap.add_argument("--limit", dest="limit", type=int)
    ap.add_argument("--grade", dest="grade", action="store_true")
    return ap.parse_args()


def main() -> None:
    """CLI entry point."""
    args = parse_args()
    summary = run_batch(args)
    if summary.get("status") == "failed":
        sys.exit(1)


if __name__ == "__main__":
    main()
