#!/usr/bin/env python3
"""
Extract student responses from two-column DOCX/PDF submission templates.
Output clean text files for review and grading.
"""
import os
import sys
import argparse
import logging
import re
from pathlib import Path
from typing import Optional, List, Tuple

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger("extract_responses")


def extract_responses_from_docx(path: str) -> Optional[str]:
    """
    Extract student responses from a two-column DOCX table.
    Assumes column 1 = question, column 2 = response.
    Returns concatenated responses or None if extraction fails.
    """
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    
    try:
        doc = Document(path)
    except Exception as e:
        logger.warning(f"Failed to read DOCX {path}: {e}")
        return None
    
    responses: List[str] = []
    
    # Find tables and extract from two-column tables
    for table in doc.tables:
        # Expect 2 columns: question | response
        if len(table.rows) == 0 or len(table.columns) < 2:
            continue
        
        # Skip header row (row 0) if it contains typical header text
        start_row = 0
        first_cell_text = (table.rows[0].cells[0].text or "").strip().lower()
        if "question" in first_cell_text or "prompt" in first_cell_text:
            start_row = 1
        
        # Extract response column (column 1)
        for row_idx in range(start_row, len(table.rows)):
            row = table.rows[row_idx]
            if len(row.cells) >= 2:
                response_cell = row.cells[1]
                response_text = (response_cell.text or "").strip()
                if response_text:
                    responses.append(response_text)
    
    if responses:
        cleaned, _ = _clean_common_boilerplate("\n\n".join(responses))
        return cleaned

    # Fallback: if no table found, return all paragraphs (less ideal)
    logger.warning(f"No two-column table found in {path}; using all paragraphs as fallback")
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    cleaned, _ = _clean_common_boilerplate("\n".join(paras) if paras else "")
    return cleaned or None


def extract_responses_from_pdf(path: str) -> Optional[str]:
    """
    Extract text from PDF. Note: PDF table extraction is complex.
    This provides basic text extraction; for structured table extraction,
    consider using pdfplumber or manual extraction.
    """
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed")
    
    try:
        reader = PdfReader(path)
    except Exception as e:
        logger.warning(f"Failed to read PDF {path}: {e}")
        return None
    
    text_parts: List[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            text_parts.append(t)

    if not text_parts:
        return None

    raw = "\n".join(text_parts)

    # Remove known multi-line instruction block spanning lines (and shared boilerplate)
    block_patterns = _shared_block_patterns()
    for bp in block_patterns:
        raw = re.sub(bp, "", raw, flags=re.IGNORECASE | re.DOTALL)

    # Try to split responses by known question prompts before line-level cleanup
    prompt_patterns = _question_prompt_patterns()
    prompt_splits = _split_by_prompts(raw, prompt_patterns)

    # --- PDF filtering: remove known template headers/instructions/prompts ---
    lines = [ln.strip() for ln in raw.splitlines()]

    # Keyword blocklist (case-insensitive) for headers/instructions
    keyword_blocklist = [
        "Note: To edit this document",
        "Week 1 Business Activity: Factors in Producing a Surfboard",
        "Use the information from Chapter",
        "Answer each question in the box below",
        "assignment drop box",
        "D2L Brightspace",
    ]

    compiled_prompts = [re.compile(p, re.IGNORECASE) for p in _question_prompt_patterns()]

    def is_template_line(ln: str) -> bool:
        low = ln.lower()
        if not low:
            return False
        # Remove page numbers or trivial headers
        if re.match(r"^\s*page\s*\d+\s*$", low):
            return True
        # Keyword blocklist
        for kw in keyword_blocklist:
            if kw.lower() in low:
                return True
        # Question prompts
        for rx in compiled_prompts:
            if rx.search(low):
                return True
        return False

    cleaned_lines = [ln for ln in lines if not is_template_line(ln)]

    # Reflow lines into paragraphs: merge line breaks, handle hyphenation
    def _reflow_lines_to_paragraphs(ls: List[str]) -> str:
        paragraphs: List[str] = []
        buf: List[str] = []
        for ln in ls:
            s = ln.strip()
            if not s:
                if buf:
                    para = " ".join(buf)
                    para = re.sub(r"\s+", " ", para).strip()
                    paragraphs.append(para)
                    buf = []
                continue
            # Join hyphenated line breaks (e.g., 'entre-
            # preneurship' -> 'entrepreneurship')
            if s.endswith("-"):
                buf.append(s[:-1])
            else:
                buf.append(s)
        if buf:
            para = " ".join(buf)
            para = re.sub(r"\s+", " ", para).strip()
            paragraphs.append(para)

        return "\n\n".join(paragraphs)

    if prompt_splits:
        cleaned_chunks = []
        for chunk in prompt_splits:
            chunk_lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
            chunk_lines = [ln for ln in chunk_lines if not is_template_line(ln)]
            chunk_text = _reflow_lines_to_paragraphs(chunk_lines).strip()
            chunk_text, _ = _clean_common_boilerplate(chunk_text)
            chunk_text = _strip_leading_bullets(chunk_text)
            if chunk_text:
                cleaned_chunks.append(chunk_text)
        if cleaned_chunks:
            return "\n\n".join(cleaned_chunks)

    out_text = _reflow_lines_to_paragraphs(cleaned_lines).strip()
    out_text, _ = _clean_common_boilerplate(out_text)
    out_text = _strip_leading_bullets(out_text)
    return out_text or None


# --- Shared boilerplate removal ---

def _shared_block_patterns() -> List[str]:
    return [
        r"Note:\s+To\s+edit\s+this\s+document.*?visible\.",
        r"Week\s+1\s+Business\s+Activity:.*?submit\s+your\s+assignment!?",
        r"Answer\s+each\s+question\s+in\s+the\s+box\s+below.*?submit\s+your\s+assignment!?",
        r"\(Box\s+will\s+expand\s+as\s+you\s+type\)",
        r"1\.\s+What\s+are\s+the\s+factors\s+of\s+production\s+needed\s+by\s+a\s+surfboard\s+manufacturer\?",
        r"2\.\s+Where\s+does\s+the\s+surfboard\s+company\s+get\s+these\s+factors\s+of\s+production\?",
        r"3\.\s+Where\s+does\s+the\s+company\s+get\s+money\s+to\s+pay\s+for\s+additional\s+resources\?",
    ]


def _question_prompt_patterns() -> List[str]:
    return [
        r"what\s+are\s+the\s+factors\s+of\s+production\s+needed\s+by\s+a\s+surfboard\s+manufacturer",
        r"where\s+does\s+the\s+surfboard\s+company\s+get\s+these\s+factors\s+of\s+production",
        r"where\s+does\s+the\s+company\s+get\s+money\s+to\s+pay\s+for\s+additional\s+resources",
    ]


def _clean_common_boilerplate(text: str) -> Tuple[str, List[str]]:
    removed: List[str] = []
    cleaned = text
    for pat in _shared_block_patterns():
        new_cleaned, n = re.subn(pat, "", cleaned, flags=re.IGNORECASE | re.DOTALL)
        if n > 0:
            removed.append(f"pattern:{pat} count:{n}")
            cleaned = new_cleaned
    # Normalize spaces while preserving paragraph breaks (\n\n)
    # Split on double newline, normalize each paragraph, then rejoin
    paras = cleaned.split("\n\n")
    normalized_paras = [re.sub(r"\s+", " ", p).strip() for p in paras if p.strip()]
    cleaned = "\n\n".join(normalized_paras)
    return cleaned, removed


def _split_by_prompts(raw_text: str, prompt_patterns: List[str]) -> List[str]:
    if not prompt_patterns:
        return []
    # Combine prompts into one regex; keep order of appearance
    combined = "|".join(f"({p})" for p in prompt_patterns)
    matches = list(re.finditer(combined, raw_text, flags=re.IGNORECASE | re.DOTALL))
    if not matches:
        return []
    segments: List[str] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        seg = raw_text[start:end].strip()
        if seg:
            segments.append(seg)
    return segments


def _strip_leading_bullets(text: str) -> str:
    """Remove stray bullet-like markers (e.g., '?', '•') at paragraph starts."""
    if not text:
        return text
    bullet_re = re.compile(r"^[\?\u2022\u2023\u2043\u25E6\-\u2013\u2014\*]+\s*")
    paras = text.split("\n\n")
    cleaned = [bullet_re.sub("", p.strip()) for p in paras]
    return "\n\n".join(cleaned).strip()


def process_directory(input_dir: str, output_dir: str, overwrite: bool = False) -> int:
    """
    Process all DOCX/PDF files in input_dir, extract responses, save to output_dir.
    Returns count of successfully extracted files.
    """
    os.makedirs(output_dir, exist_ok=True)
    
    count = 0
    for root, _, files in os.walk(input_dir):
        for name in sorted(files):
            full_path = os.path.join(root, name)
            lower_name = name.lower()
            
            # Skip already-extracted files unless overwrite requested
            if ("_extracted" in lower_name or "_clean" in lower_name) and not overwrite:
                continue
            
            if lower_name.endswith(".docx"):
                responses = extract_responses_from_docx(full_path)
                source_type = "DOCX"
            elif lower_name.endswith(".pdf"):
                responses = extract_responses_from_pdf(full_path)
                source_type = "PDF"
            else:
                continue
            
            if responses:
                # Output filename: base_name_extracted.txt
                base_name = os.path.splitext(name)[0]
                output_name = f"{base_name}_extracted.txt"
                output_path = os.path.join(output_dir, output_name)
                
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(responses)
                
                logger.info(f"✓ Extracted {source_type}: {name} → {output_name}")
                count += 1
            else:
                logger.warning(f"✗ Failed to extract responses from {source_type}: {name}")
    
    return count


def main():
    ap = argparse.ArgumentParser(description="Extract student responses from submission templates (DOCX/PDF)")
    ap.add_argument("input_dir", help="Directory containing raw DOCX/PDF submission files")
    ap.add_argument("output_dir", default=".", nargs="?", help="Directory to save extracted .txt files (default: current dir)")
    ap.add_argument("--overwrite", action="store_true", help="Overwrite existing _extracted files if present")
    
    args = ap.parse_args()
    
    if not os.path.isdir(args.input_dir):
        logger.error(f"Input directory not found: {args.input_dir}")
        sys.exit(1)
    
    logger.info(f"Extracting responses from {args.input_dir}...")
    count = process_directory(args.input_dir, args.output_dir, overwrite=args.overwrite)
    
    logger.info(f"Done. Extracted {count} file(s) to {args.output_dir}/")
    logger.info("Review the extracted files, then run batch grading on them.")


if __name__ == "__main__":
    main()
